#!/usr/bin/env python3
"""Render an EBA protokoll Markdown file with the matching official template.

Used by all five format skills. The Markdown is treated as an in-memory
intermediate and is NOT preserved in the user-facing output. The deliverables
follow the original QMG source format: Word templates produce .docx plus .pdf,
while Excel templates produce .xlsx.

Windows 11 + MS Word is the primary production path. The script self-bootstraps
its Python dependencies in the current user environment, then exports PDF with
MS Word COM via pywin32. LibreOffice and macOS Pages are fallback converters.

The script auto-detects the format from the Markdown header if --format is not
given:
    "# Gesprächsnotiz"          -> gespraechsnotiz
    "# Protokoll" + "Gesprächsinhalt" + "Frist" header -> protokoll-einfach
    "# Protokoll" + "Besprechungsthemen" + "D/K"       -> protokoll-lp1-4 / protokoll-lp5 / protokoll-bim

Output files are written next to <markdown_path>:
    <basename>.docx + <basename>.pdf (Word-origin formats)
    <basename>.xlsx (Excel-origin formats)

The Markdown intermediate at <markdown_path> is always removed on success.

Exit codes:
    0 success (the format-specific deliverables were written)
    2 markdown could not be parsed
    3 docx generation failed
    4 required PDF generation failed
    5 XLSX generation failed
    6 Markdown intermediate could not be removed
"""

from __future__ import annotations

import argparse
from copy import copy, deepcopy
import json
import re
import shutil
import site
import subprocess
import sys
from dataclasses import dataclass, field
from pathlib import Path


def _in_virtualenv() -> bool:
    return sys.prefix != getattr(sys, "base_prefix", sys.prefix)


def _run_pip_install(args: list[str]) -> bool:
    cmd = [sys.executable, "-m", "pip", "install", "--disable-pip-version-check"]
    if not _in_virtualenv():
        cmd.append("--user")
    cmd.extend(args)
    pip_error = ""
    try:
        result = subprocess.run(
            cmd,
            text=True,
            capture_output=True,
            timeout=300,
        )
        if result.returncode == 0:
            return True
        pip_error = result.stderr.strip()
    except Exception as exc:
        pip_error = str(exc)

    if _run_uv_install(args):
        return True
    if pip_error:
        sys.stderr.write(pip_error + "\n")
    return False


def _run_uv_install(args: list[str]) -> bool:
    """Fallback for lean Python installs that have uv but no pip/ensurepip."""
    uv = shutil.which("uv")
    if not uv:
        return False
    try:
        if _in_virtualenv():
            site_packages = site.getsitepackages()
            target = site_packages[0] if site_packages else site.getusersitepackages()
        else:
            target = site.getusersitepackages()
        Path(target).mkdir(parents=True, exist_ok=True)
        result = subprocess.run(
            [
                uv,
                "pip",
                "install",
                "--python",
                sys.executable,
                "--target",
                target,
                *args,
            ],
            text=True,
            capture_output=True,
            timeout=300,
        )
        if result.returncode == 0:
            return True
        sys.stderr.write(result.stderr.strip() + "\n")
    except Exception as exc:
        sys.stderr.write(f"uv dependency bootstrap failed: {exc}\n")
    return False


def _ensure_pip() -> None:
    check = subprocess.run(
        [sys.executable, "-m", "pip", "--version"],
        text=True,
        capture_output=True,
    )
    if check.returncode == 0:
        return
    subprocess.run([sys.executable, "-m", "ensurepip", "--upgrade"], check=False)


def _bootstrap_requirements() -> bool:
    """Install renderer dependencies without user interaction.

    This is intentionally quiet and user-local. The Codex user should not
    need to know what python-docx, openpyxl, or pywin32 are.
    """
    requirements = Path(__file__).resolve().parent / "requirements.txt"
    if not requirements.exists():
        sys.stderr.write(f"Missing requirements file: {requirements}\n")
        return False
    _ensure_pip()
    sys.stderr.write("Renderer dependency bootstrap: installing required packages...\n")
    return _run_pip_install(["--upgrade", "-r", str(requirements)])


def _bootstrap_package(package: str) -> bool:
    _ensure_pip()
    sys.stderr.write(f"Renderer dependency bootstrap: installing {package}...\n")
    return _run_pip_install(["--upgrade", package])


try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from openpyxl import load_workbook
    from openpyxl.cell.cell import MergedCell
    from openpyxl.utils import get_column_letter
    from openpyxl.worksheet.cell_range import CellRange
    from openpyxl.worksheet.table import TableColumn
except ImportError:
    if not _bootstrap_requirements():
        sys.stderr.write(
            "render_protokoll.py: could not install required DOCX dependencies.\n"
        )
        sys.exit(3)
    try:
        from docx import Document
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from openpyxl import load_workbook
        from openpyxl.cell.cell import MergedCell
        from openpyxl.utils import get_column_letter
        from openpyxl.worksheet.cell_range import CellRange
        from openpyxl.worksheet.table import TableColumn
    except ImportError as exc:
        sys.stderr.write(
            "render_protokoll.py: DOCX dependencies are still unavailable after "
            f"bootstrap: {exc}\n"
        )
        sys.exit(3)


QMG_TEMPLATE_DIR = Path(__file__).resolve().parents[1] / "references" / "templates" / "qmg"
GESPRAECHSNOTIZ_TEMPLATE = QMG_TEMPLATE_DIR / "QMG-024-141_ORG-GESPRAECHSNOTIZ_230202-D.docx"
PROTOKOLL_EINFACH_TEMPLATE = QMG_TEMPLATE_DIR / "QMG-024-141_ORG-PK-LP1-4-MA_230227-A.docx"
TRACKING_WORD_TEMPLATE = QMG_TEMPLATE_DIR / "QMG-024-141_ORG-PK-LP5-MA_230202-B.docx"
PROTOKOLL_EINFACH_EXCEL_TEMPLATE = QMG_TEMPLATE_DIR / "QMG-024-141_ORG-PK-LP1-4-EXCEL-MA_240920-A.xlsx"
TRACKING_EXCEL_TEMPLATE = QMG_TEMPLATE_DIR / "QMG-024-141_ORG-PK-EXCEL-MA_240926-C.xlsx"


# ─── Markdown parsing ──────────────────────────────────────────────────────


@dataclass
class MdSection:
    """One ## section of the protokoll: heading + raw lines (no leading ##)."""

    heading: str
    lines: list[str] = field(default_factory=list)


@dataclass
class ParsedMd:
    """Structured view of an EBA protokoll Markdown file."""

    title: str = ""
    subtitle: str = ""  # the italic line right under the title (if any)
    header_tables: list[list[list[str]]] = field(default_factory=list)
    notice: str = ""  # the > Hinweis / > Vorbemerkung blockquote
    sections: list[MdSection] = field(default_factory=list)
    detected_format: str = "unknown"


def _strip_md_inline(s: str) -> str:
    """Remove a *small* set of inline Markdown tokens (bold, italics) so the
    resulting text reads as plain Word body text. We deliberately keep the
    behaviour conservative — anything we don't recognise stays as-is."""
    s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)
    s = re.sub(r"__(.+?)__", r"\1", s)
    s = re.sub(r"(?<![A-Za-z0-9])\*(.+?)\*(?![A-Za-z0-9])", r"\1", s)
    s = re.sub(r"(?<![A-Za-z0-9])_(.+?)_(?![A-Za-z0-9])", r"\1", s)
    return s


def _parse_md_table(block: list[str]) -> list[list[str]]:
    """Parse a Markdown pipe-table block into a list of rows of cell strings.
    The separator row (---|---) is skipped."""
    rows: list[list[str]] = []
    for raw in block:
        line = raw.strip()
        if not line.startswith("|") or not line.endswith("|"):
            continue
        cells = [c.strip() for c in line[1:-1].split("|")]
        if all(re.fullmatch(r":?-+:?", c or "-") for c in cells):
            continue  # separator row
        rows.append([_strip_md_inline(c) for c in cells])
    return rows


def parse_protokoll_md(md_text: str) -> ParsedMd:
    """Walk the markdown line by line, collecting:
       - title (# H1)
       - subtitle (italic line directly under H1)
       - header tables (the tables before the first ## heading)
       - notice block (>... immediately after header tables)
       - body sections (## Heading + content)
    """
    parsed = ParsedMd()
    lines = md_text.splitlines()
    i = 0

    # --- title ---
    while i < len(lines) and not lines[i].startswith("# "):
        i += 1
    if i < len(lines):
        parsed.title = lines[i][2:].strip()
        i += 1

    # --- subtitle: skip blank lines, then capture an italic line if present ---
    while i < len(lines) and not lines[i].strip():
        i += 1
    if i < len(lines):
        m = re.match(r"^_(.+)_\s*$", lines[i].strip())
        if m:
            parsed.subtitle = m.group(1).strip()
            i += 1

    # --- collect everything until first ## as header / notice content ---
    pre_section: list[str] = []
    while i < len(lines) and not lines[i].startswith("## "):
        pre_section.append(lines[i])
        i += 1

    # split pre-section into tables / blockquotes / other
    j = 0
    cur_table: list[str] = []
    cur_quote: list[str] = []
    while j < len(pre_section):
        line = pre_section[j]
        if line.lstrip().startswith("|"):
            cur_table.append(line)
        else:
            if cur_table:
                parsed.header_tables.append(_parse_md_table(cur_table))
                cur_table = []
            if line.lstrip().startswith(">"):
                cur_quote.append(re.sub(r"^>\s?", "", line.strip()))
            elif cur_quote and not line.strip():
                pass  # allow inline blank lines inside quote
            elif cur_quote:
                parsed.notice = " ".join(cur_quote).strip()
                cur_quote = []
        j += 1
    if cur_table:
        parsed.header_tables.append(_parse_md_table(cur_table))
    if cur_quote and not parsed.notice:
        parsed.notice = " ".join(cur_quote).strip()

    # --- body sections ---
    cur_section: MdSection | None = None
    while i < len(lines):
        line = lines[i]
        if line.startswith("## "):
            if cur_section:
                parsed.sections.append(cur_section)
            cur_section = MdSection(heading=line[3:].strip())
        else:
            if cur_section is not None:
                cur_section.lines.append(line)
        i += 1
    if cur_section:
        parsed.sections.append(cur_section)

    parsed.detected_format = _detect_format(parsed)
    return parsed


def _detect_format(parsed: ParsedMd) -> str:
    if parsed.title.startswith("Gesprächsnotiz"):
        return "gespraechsnotiz"
    if parsed.title.startswith("Protokoll"):
        all_text = "\n".join(
            line for s in parsed.sections for line in [s.heading] + s.lines
        )
        if "D/K" in all_text and "Besprechungsthemen" in all_text:
            head = "\n".join(
                line for s in parsed.sections[:2] for line in [s.heading] + s.lines[:12]
            )
            if re.search(
                r"\b(BIM-Koordination|BIM-Jour-Fixe|BIM-JF|Koordinations-JF|BIM-Termin)\b",
                head,
                re.IGNORECASE,
            ):
                return "protokoll-bim"
            if re.search(
                r"\b(Baubesprechung|Baustelle|Mängel|M-\d+|Witterung|Rohbau|Polier|"
                r"Gewerk|Abnahme|Bemusterung|Objektüberwachung)\b",
                head + "\n" + all_text,
                re.IGNORECASE,
            ):
                return "protokoll-lp5"
            return "protokoll-lp1-4"
        if "Gesprächsinhalt" in all_text:
            return "protokoll-einfach"
    return "unknown"


def _first_table_from_section(parsed: ParsedMd, heading: str) -> list[list[str]]:
    for section in parsed.sections:
        if section.heading.strip().lower() != heading.lower():
            continue
        block: list[str] = []
        for line in section.lines:
            if line.lstrip().startswith("|"):
                block.append(line)
            elif block:
                break
        return _parse_md_table(block)
    return []


def _tables_from_lines(lines: list[str]) -> list[list[list[str]]]:
    tables: list[list[list[str]]] = []
    block: list[str] = []
    for line in lines:
        if line.lstrip().startswith("|"):
            block.append(line)
        elif block:
            tables.append(_parse_md_table(block))
            block = []
    if block:
        tables.append(_parse_md_table(block))
    return [table for table in tables if table]


def _section_by_prefix(parsed: ParsedMd, prefix: str) -> MdSection | None:
    prefix = prefix.lower()
    for section in parsed.sections:
        if section.heading.strip().lower().startswith(prefix):
            return section
    return None


def _quote_text_from_lines(lines: list[str]) -> str:
    quote: list[str] = []
    for line in lines:
        stripped = line.lstrip()
        if stripped.startswith(">"):
            quote.append(re.sub(r"^>\s?", "", stripped).strip())
        elif quote and stripped:
            break
    return _strip_md_inline(" ".join(quote).strip())


def _first_subheading(lines: list[str]) -> str:
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("### "):
            return _strip_md_inline(stripped[4:].strip())
    return ""


def _kv_from_table(table: list[list[str]]) -> dict[str, str]:
    values: dict[str, str] = {}
    for row in table:
        if len(row) >= 2:
            values[row[0].strip()] = row[1].strip()
    return values


def _pick(values: dict[str, str], *keys: str) -> str:
    normalized = {
        re.sub(r"[\s.\-_/]+", "", key).lower(): value for key, value in values.items()
    }
    for key in keys:
        value = normalized.get(re.sub(r"[\s.\-_/]+", "", key).lower())
        if value is not None:
            return value
    return ""


def _actual_tcs(row_or_tr) -> list:
    tr = getattr(row_or_tr, "_tr", row_or_tr)
    return list(tr.findall(qn("w:tc")))


def _first_rpr(p):
    for r in p.findall(qn("w:r")):
        rpr = r.find(qn("w:rPr"))
        if rpr is not None:
            return deepcopy(rpr)
    return None


def _set_tc_text(tc, text: str) -> None:
    """Replace visible text in a template table cell while keeping cell and
    paragraph formatting. This avoids flattening the QMG grid styling.

    The QMG headers use REF fields whose targets are bookmarks embedded in
    several of these cells. Keep the bookmark boundaries around the new text;
    deleting them makes Word/LibreOffice render "Reference source not found"
    in continuation-page headers.
    """
    text = _strip_md_inline(text or "")
    tc_el = getattr(tc, "_tc", tc)
    paragraphs = tc_el.findall(qn("w:p"))
    if not paragraphs:
        paragraphs = [OxmlElement("w:p")]
        tc_el.append(paragraphs[0])
    p = paragraphs[0]
    rpr_template = _first_rpr(p)
    bookmark_starts = [
        deepcopy(child)
        for paragraph in paragraphs
        for child in paragraph
        if child.tag == qn("w:bookmarkStart")
    ]
    bookmark_ends = [
        deepcopy(child)
        for paragraph in paragraphs
        for child in paragraph
        if child.tag == qn("w:bookmarkEnd")
    ]
    ppr = p.find(qn("w:pPr"))
    for child in list(p):
        if child is not ppr:
            p.remove(child)
    for extra in paragraphs[1:]:
        tc_el.remove(extra)

    r = OxmlElement("w:r")
    if rpr_template is not None:
        r.append(rpr_template)
    parts = text.split("\n") if text else [""]
    for idx, part in enumerate(parts):
        if idx:
            r.append(OxmlElement("w:br"))
        t = OxmlElement("w:t")
        if part.startswith(" ") or part.endswith(" "):
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t.text = part
        r.append(t)
    for bookmark in bookmark_starts:
        p.append(bookmark)
    p.append(r)
    for bookmark in bookmark_ends:
        p.append(bookmark)


def _replace_visible_text(element, replacements: dict[str, str]) -> None:
    for text_node in element.iter(qn("w:t")):
        if not text_node.text:
            continue
        new_text = text_node.text
        for old, new in replacements.items():
            new_text = new_text.replace(old, new)
        text_node.text = new_text


def _prepare_qmg_fields(doc) -> None:
    """Make copied QMG header/footer fields portable and refreshable.

    One source template stores the PrjNr REF switch before its bookmark name,
    which Word accepts but LibreOffice does not. SECTIONPAGES is also cached
    against the untrimmed multi-section source document. The rendered files
    have one section, so NUMPAGES is the equivalent field and yields the real
    output page total in both converters.
    """
    for part in doc.part.related_parts.values():
        if not (
            part.partname.startswith("/word/header")
            or part.partname.startswith("/word/footer")
        ):
            continue
        for instruction in part._element.iter(qn("w:instrText")):
            value = instruction.text or ""
            value = re.sub(
                r"(\bREF\s+)\\\*\s+CHARFORMAT\s+([A-Za-z0-9_]+)",
                r"\1\2 \\* CHARFORMAT",
                value,
            )
            instruction.text = value.replace("SECTIONPAGES", "NUMPAGES")

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def _delete_rows_after(table, keep_last_index: int) -> None:
    for row in list(table.rows)[keep_last_index + 1 :]:
        row._tr.getparent().remove(row._tr)


def _append_row_from_template(table, template_tr):
    new_tr = deepcopy(template_tr)
    table._tbl.append(new_tr)
    return new_tr


def _fill_row_cells(tr, values: list[str], cell_indices: list[int]) -> None:
    cells = _actual_tcs(tr)
    for value, cell_idx in zip(values, cell_indices):
        if cell_idx < len(cells):
            _set_tc_text(cells[cell_idx], value)


def _fill_grid_table(table, rows: list[list[str]], *, start_row: int, cell_indices: list[int]) -> None:
    templates = [deepcopy(row._tr) for row in table.rows[start_row:]]
    fallback = templates[0] if templates else deepcopy(table.rows[-1]._tr)
    _delete_rows_after(table, start_row - 1)
    for values in rows:
        tr = _append_row_from_template(table, fallback)
        _fill_row_cells(tr, values, cell_indices)


def _replace_row_range(
    table,
    start_idx: int,
    end_idx: int,
    rows: list[list[str]],
    *,
    cell_indices: list[int],
    template_idx: int | None = None,
) -> None:
    existing_rows = list(table.rows)
    if start_idx >= len(existing_rows) or end_idx < start_idx:
        return
    end_idx = min(end_idx, len(existing_rows) - 1)
    template_source = existing_rows[template_idx if template_idx is not None else start_idx]._tr
    template_tr = deepcopy(template_source)
    anchor = existing_rows[end_idx + 1]._tr if end_idx + 1 < len(existing_rows) else None
    for row in reversed(existing_rows[start_idx : end_idx + 1]):
        row._tr.getparent().remove(row._tr)
    for values in rows:
        new_tr = deepcopy(template_tr)
        _fill_row_cells(new_tr, values, cell_indices)
        if anchor is not None:
            anchor.addprevious(new_tr)
        else:
            table._tbl.append(new_tr)


def _fill_gespraechsinhalt_table(table, rows: list[list[str]]) -> None:
    template_main = deepcopy(table.rows[1]._tr)
    template_sub = deepcopy(table.rows[2]._tr if len(table.rows) > 2 else table.rows[1]._tr)
    _delete_rows_after(table, 0)
    for row in rows:
        topic = row[0] if len(row) > 0 else ""
        desc = row[1] if len(row) > 1 else ""
        owner = row[2] if len(row) > 2 else ""
        template = template_main if re.fullmatch(r"Thema\s+\d+", topic.strip()) else template_sub
        tr = _append_row_from_template(table, template)
        cells = _actual_tcs(tr)
        for value, cell_idx in [(topic, 0), (desc, 2), (owner, 4)]:
            if cell_idx < len(cells):
                _set_tc_text(cells[cell_idx], value)


def _trim_qmg_template_body(doc, *, keep_tables: int = 4) -> None:
    """Keep only the official front document and discard QMG help/internal pages."""
    body = doc._body._body
    children = list(body)
    first_sect_pr = None
    for child in children:
        if child.tag == qn("w:p"):
            ppr = child.find(qn("w:pPr"))
            if ppr is not None:
                sect_pr = ppr.find(qn("w:sectPr"))
                if sect_pr is not None:
                    first_sect_pr = deepcopy(sect_pr)
                    break
    if first_sect_pr is None:
        first_sect_pr = deepcopy(doc.sections[0]._sectPr)

    table_count = 0
    cut_index = len(children)
    for idx, child in enumerate(children):
        if child.tag == qn("w:tbl"):
            table_count += 1
            if table_count == keep_tables:
                cut_index = idx + 1
                break
    for child in children[cut_index:]:
        body.remove(child)

    for child in list(body):
        if child.tag == qn("w:sectPr"):
            body.remove(child)
    body.append(first_sect_pr)


def _render_simple_word_template(
    parsed: ParsedMd,
    out_path: Path,
    *,
    template_path: Path,
    default_title: str,
) -> None:
    if not template_path.exists():
        raise FileNotFoundError(f"Missing QMG template: {template_path}")

    doc = Document(str(template_path))
    _trim_qmg_template_body(doc, keep_tables=4)

    project = _kv_from_table(parsed.header_tables[0] if parsed.header_tables else [])
    meta_table = parsed.header_tables[1] if len(parsed.header_tables) > 1 else []
    meta = dict(zip(meta_table[0], meta_table[1])) if len(meta_table) >= 2 else {}

    project_name = project.get("Projektname", "")
    project_number = project.get("Projekt-Nummer", "")
    project_desc = project.get("Projekt-Beschreibung", "")
    subtitle = parsed.subtitle or project_desc

    main = doc.tables[0]
    _set_tc_text(_actual_tcs(main.rows[0])[0], parsed.title or default_title)
    _set_tc_text(_actual_tcs(main.rows[1])[0], subtitle)
    _set_tc_text(_actual_tcs(main.rows[2])[2], project_name)
    _set_tc_text(_actual_tcs(main.rows[3])[2], project_number)
    _set_tc_text(_actual_tcs(main.rows[4])[2], project_desc)
    row = _actual_tcs(main.rows[7])
    for key, idx in [("Ort", 0), ("Gesprächsdatum", 2), ("Erstelldatum", 4), ("Ersteller", 6)]:
        if idx < len(row):
            _set_tc_text(row[idx], meta.get(key, ""))

    teilnehmer = _first_table_from_section(parsed, "Teilnehmer")
    teilnehmer_rows = [r for r in teilnehmer[1:] if any(c.strip() for c in r)]
    _fill_grid_table(doc.tables[1], teilnehmer_rows, start_row=2, cell_indices=[0, 2, 4, 6])

    verteiler = _first_table_from_section(parsed, "Verteiler")
    verteiler_rows = [r for r in verteiler[1:] if any(c.strip() for c in r)]
    _fill_grid_table(doc.tables[2], verteiler_rows, start_row=2, cell_indices=[0, 2, 6])

    inhalt = _first_table_from_section(parsed, "Gesprächsinhalt")
    inhalt_rows = [r for r in inhalt[1:] if any(c.strip() for c in r)]
    _fill_gespraechsinhalt_table(doc.tables[3], inhalt_rows)

    replacements = {
        "_Projektnummer einsetzen_": project_number,
        "_Projektname einsetzen_": project_name,
        "_kurze Beschreibung zum Dokument/Übergeordnetes Thema einsetzen_": subtitle,
        "_Projektbeschreibung_": project_desc,
        "_Ort_": meta.get("Ort", ""),
        "_tt.mm.jj_": meta.get("Gesprächsdatum", ""),
        "_Ersteller_": meta.get("Ersteller", ""),
    }
    _replace_visible_text(doc._element, replacements)
    for part in doc.part.related_parts.values():
        if part.partname.startswith("/word/header") or part.partname.startswith("/word/footer"):
            _replace_visible_text(part._element, replacements)
    _prepare_qmg_fields(doc)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def _render_gespraechsnotiz_template(parsed: ParsedMd, out_path: Path) -> None:
    _render_simple_word_template(
        parsed,
        out_path,
        template_path=GESPRAECHSNOTIZ_TEMPLATE,
        default_title="Gesprächsnotiz",
    )


def _render_protokoll_einfach_template(parsed: ParsedMd, out_path: Path) -> None:
    _render_simple_word_template(
        parsed,
        out_path,
        template_path=PROTOKOLL_EINFACH_TEMPLATE,
        default_title="Protokoll",
    )


def _clean_rows(rows: list[list[str]]) -> list[list[str]]:
    return [row for row in rows if any(cell.strip() for cell in row)]


def _pad_row(row: list[str], length: int) -> list[str]:
    return (row + [""] * length)[:length]


def _is_empty_marker(value: str) -> bool:
    return value.strip() in {"", "-", "–", "—"}


def _tracking_data_rows(parsed: ParsedMd, heading: str) -> list[list[str]]:
    table = _first_table_from_section(parsed, heading)
    return _clean_rows(table[1:] if table else [])


def _tracking_header(parsed: ParsedMd) -> dict[str, str]:
    intro = parsed.sections[0] if parsed.sections else MdSection("")
    intro_tables = _tables_from_lines(intro.lines)
    project = _kv_from_table(intro_tables[0] if intro_tables else [])
    meta = _kv_from_table(intro_tables[1] if len(intro_tables) > 1 else [])
    meeting_no_match = re.search(r"Nr\.\s*([A-Za-z0-9._-]+)", intro.heading)
    meeting_no = meeting_no_match.group(1) if meeting_no_match else ""
    return {
        "meeting_no": meeting_no,
        "meeting_topic": _first_subheading(intro.lines) or parsed.subtitle or "Besprechung",
        "notice": _quote_text_from_lines(intro.lines),
        "project_number": _pick(project, "Projekt-Nr.", "Projekt-Nr", "Projekt-Nummer"),
        "project_name": _pick(project, "Projekt-Name", "Projektname"),
        "ort": _pick(meta, "Ort"),
        "datum": _pick(meta, "Datum", "Gesprächsdatum"),
        "zeit": _pick(meta, "Zeit"),
    }


def _render_tracking_template(parsed: ParsedMd, out_path: Path) -> None:
    if not TRACKING_WORD_TEMPLATE.exists():
        raise FileNotFoundError(f"Missing QMG template: {TRACKING_WORD_TEMPLATE}")

    doc = Document(str(TRACKING_WORD_TEMPLATE))
    _trim_qmg_template_body(doc, keep_tables=4)

    header = _tracking_header(parsed)
    meeting_no = header["meeting_no"]
    meeting_topic = header["meeting_topic"]
    notice = header["notice"]
    project_number = header["project_number"]
    project_name = header["project_name"]
    ort = header["ort"]
    datum = header["datum"]
    zeit = header["zeit"]

    main = doc.tables[0]
    _set_tc_text(_actual_tcs(main.rows[2])[0], f"zur Besprechung Nr. {meeting_no}".strip())
    _set_tc_text(_actual_tcs(main.rows[5])[1], meeting_topic)
    _set_tc_text(_actual_tcs(main.rows[6])[1], project_number)
    if notice:
        _set_tc_text(_actual_tcs(main.rows[6])[3], notice)
    _set_tc_text(_actual_tcs(main.rows[7])[1], project_name)
    _set_tc_text(_actual_tcs(main.rows[9])[1], ort)
    _set_tc_text(_actual_tcs(main.rows[9])[4], datum)
    _set_tc_text(_actual_tcs(main.rows[10])[4], zeit)

    unterlagen = _tracking_data_rows(parsed, "Besprochene Unterlagen")
    _replace_row_range(
        main,
        28,
        30,
        unterlagen,
        cell_indices=[0, 1, 2, 3, 4],
        template_idx=28,
    )

    teilnehmer = _tracking_data_rows(parsed, "Teilnehmer")
    _replace_row_range(
        main,
        14,
        23,
        teilnehmer,
        cell_indices=[0, 1, 2, 3, 4, 5, 6],
        template_idx=14,
    )

    themen = _tracking_data_rows(parsed, "Besprechungsthemen")
    theme_table = doc.tables[1]
    category_template = deepcopy(theme_table.rows[3]._tr)
    topic_template = deepcopy(theme_table.rows[4]._tr)
    _delete_rows_after(theme_table, 1)
    for row in themen:
        padded = row + [""] * (7 - len(row))
        is_category = _is_empty_marker(padded[1]) and _is_empty_marker(padded[2])
        tr = _append_row_from_template(
            theme_table,
            category_template if is_category else topic_template,
        )
        if is_category:
            _fill_row_cells(tr, [padded[0], padded[3]], [0, 1])
        else:
            _fill_row_cells(tr, padded[:7], [0, 1, 2, 3, 4, 5, 6])

    termine_section = _section_by_prefix(parsed, "Termine")
    termine = []
    if termine_section is not None:
        tables = _tables_from_lines(termine_section.lines)
        termine = _clean_rows(tables[0][1:] if tables else [])
    _replace_row_range(
        doc.tables[2],
        2,
        4,
        termine,
        cell_indices=[0, 1, 3, 4, 5],
        template_idx=2,
    )

    tail = doc.tables[3]
    aufstell = _tracking_data_rows(parsed, "Aufstellvermerk zum Dokument")
    aufstell_by_role = {row[0].strip().lower(): row for row in aufstell if row}
    for role, row_idx in [("ersteller", 2), ("geprüft", 3)]:
        row = aufstell_by_role.get(role, [])
        cells = _actual_tcs(tail.rows[row_idx])
        if len(cells) > 1:
            _set_tc_text(cells[1], row[1] if len(row) > 1 else "")
        if len(cells) > 5:
            _set_tc_text(cells[5], row[2] if len(row) > 2 else "")

    anmerkungen = _tracking_data_rows(parsed, "Nachträgliche Anmerkungen zum Dokument")
    anmerkung = anmerkungen[0] if anmerkungen else []
    row6 = _actual_tcs(tail.rows[6])
    row7 = _actual_tcs(tail.rows[7])
    if len(row6) > 1:
        _set_tc_text(row6[1], anmerkung[0] if len(anmerkung) > 0 else "- - -")
    if len(row6) > 5:
        _set_tc_text(row6[5], anmerkung[1] if len(anmerkung) > 1 else "- - -")
    if len(row7) > 1:
        _set_tc_text(row7[1], anmerkung[2] if len(anmerkung) > 2 else "- - -")

    kennz = _tracking_data_rows(parsed, "Kennzeichnungen im Dokument")
    _replace_row_range(
        tail,
        17,
        19,
        kennz,
        cell_indices=[0, 1, 3],
        template_idx=17,
    )

    anlagen = _tracking_data_rows(parsed, "Anlagen")
    _replace_row_range(
        tail,
        11,
        13,
        anlagen,
        cell_indices=[0, 1, 2, 4],
        template_idx=11,
    )

    replacements = {
        "_Prj.-Nr._": project_number,
        "_Prj.-Name_": project_name,
        "_Besprechungsthema_": meeting_topic,
        "zur Besprechung Nr. XX": f"zur Besprechung Nr. {meeting_no}".strip(),
        "_Ort_": ort,
        "TT.MM.JJ": datum,
        "0.00 – 0.00 Uhr": zeit,
    }
    _replace_visible_text(doc._element, replacements)
    for part in doc.part.related_parts.values():
        if part.partname.startswith("/word/header") or part.partname.startswith("/word/footer"):
            _replace_visible_text(part._element, replacements)
    _prepare_qmg_fields(doc)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


# ─── XLSX rendering from official QMG templates ───────────────────────────


def _is_tracking_excel_format(parsed: ParsedMd) -> bool:
    return parsed.detected_format in {
        "protokoll-bim",
        "protokoll-lp1-4-excel",
    }


def _is_simple_excel_format(parsed: ParsedMd) -> bool:
    return parsed.detected_format == "protokoll-einfach-excel"


def _is_xlsx_only_format(parsed: ParsedMd) -> bool:
    return _is_simple_excel_format(parsed) or _is_tracking_excel_format(parsed)


def _excel_safe(value: str | None) -> str:
    return _strip_md_inline(value or "")


def _excel_set(cell, value: str | None) -> None:
    if isinstance(cell, MergedCell):
        return
    cell.value = _excel_safe(value)


def _excel_set_raw(cell, value) -> None:
    if isinstance(cell, MergedCell):
        return
    cell.value = value


def _tracking_meeting_number_value(meeting_no: str):
    stripped = (meeting_no or "").strip()
    if re.fullmatch(r"\d+", stripped):
        return int(stripped)
    return stripped


def _excel_clear_row(ws, row_idx: int, *, max_col: int) -> None:
    for col_idx in range(1, max_col + 1):
        cell = ws.cell(row_idx, col_idx)
        if not isinstance(cell, MergedCell):
            cell.value = None


def _excel_copy_cell(src, dst) -> None:
    if isinstance(dst, MergedCell):
        return
    if src.has_style:
        dst._style = copy(src._style)
    if src.number_format:
        dst.number_format = src.number_format
    if src.comment:
        dst.comment = copy(src.comment)
    if src.hyperlink:
        dst._hyperlink = copy(src.hyperlink)


def _excel_copy_row_template(ws, source_row: int, target_row: int, *, max_col: int) -> None:
    ws.row_dimensions[target_row].height = ws.row_dimensions[source_row].height
    for col_idx in range(1, max_col + 1):
        _excel_copy_cell(ws.cell(source_row, col_idx), ws.cell(target_row, col_idx))

    for merged in list(ws.merged_cells.ranges):
        if merged.min_row == merged.max_row == source_row:
            target_ref = (
                f"{get_column_letter(merged.min_col)}{target_row}:"
                f"{get_column_letter(merged.max_col)}{target_row}"
            )
            if target_ref not in {str(rng) for rng in ws.merged_cells.ranges}:
                try:
                    ws.merge_cells(target_ref)
                except ValueError:
                    pass


def _excel_insert_rows(ws, idx: int, amount: int) -> None:
    merged_ranges = [CellRange(str(rng)) for rng in ws.merged_cells.ranges]
    for merged in list(ws.merged_cells.ranges):
        ws.unmerge_cells(str(merged))

    ws.insert_rows(idx, amount)

    for merged in merged_ranges:
        if merged.min_row >= idx:
            merged.shift(row_shift=amount)
        elif merged.max_row >= idx:
            merged.max_row += amount
        ws.merge_cells(str(merged))


def _excel_ensure_capacity(
    ws,
    *,
    start_row: int,
    capacity: int,
    needed: int,
    template_row: int,
    max_col: int,
) -> int:
    extra = max(0, needed - capacity)
    if extra:
        insert_at = start_row + capacity
        _excel_insert_rows(ws, insert_at, extra)
        for row_idx in range(insert_at, insert_at + extra):
            _excel_copy_row_template(ws, template_row, row_idx, max_col=max_col)
    return extra


def _excel_put_row(
    ws,
    row_idx: int,
    values: list[str],
    cell_indices: list[int],
    *,
    template_row: int,
    max_col: int,
) -> None:
    if row_idx != template_row:
        _excel_copy_row_template(ws, template_row, row_idx, max_col=max_col)
    _excel_clear_row(ws, row_idx, max_col=max_col)
    for value, col_idx in zip(values, cell_indices):
        _excel_set(ws.cell(row_idx, col_idx), value)


def _excel_fill_block(
    ws,
    *,
    start_row: int,
    capacity: int,
    rows: list[list[str]],
    cell_indices: list[int],
    template_row: int,
    max_col: int,
) -> int:
    extra = _excel_ensure_capacity(
        ws,
        start_row=start_row,
        capacity=capacity,
        needed=len(rows),
        template_row=template_row,
        max_col=max_col,
    )
    total_rows = max(capacity + extra, len(rows))
    for offset in range(total_rows):
        target_row = start_row + offset
        if offset < len(rows):
            _excel_put_row(
                ws,
                target_row,
                rows[offset],
                cell_indices,
                template_row=template_row,
                max_col=max_col,
            )
        else:
            _excel_clear_row(ws, target_row, max_col=max_col)
    return extra


def _set_print_area(ws, last_col: str, last_row: int) -> None:
    ws.print_area = f"A1:{last_col}{last_row}"


def _simple_header(parsed: ParsedMd) -> dict[str, str]:
    project = _kv_from_table(parsed.header_tables[0] if parsed.header_tables else [])
    meta_table = parsed.header_tables[1] if len(parsed.header_tables) > 1 else []
    meta = dict(zip(meta_table[0], meta_table[1])) if len(meta_table) >= 2 else {}
    project_desc = project.get("Projekt-Beschreibung", "")
    return {
        "title": parsed.title or "Protokoll",
        "subtitle": parsed.subtitle or project_desc or "Besprechung",
        "project_name": project.get("Projektname", ""),
        "project_number": project.get("Projekt-Nummer", ""),
        "project_desc": project_desc,
        "ort": meta.get("Ort", ""),
        "gespraechsdatum": meta.get("Gesprächsdatum", ""),
        "erstelldatum": meta.get("Erstelldatum", ""),
        "ersteller": meta.get("Ersteller", ""),
        "notice": parsed.notice,
    }


def _split_owner_deadline(value: str) -> tuple[str, str]:
    if " / " in value:
        owner, deadline = value.split(" / ", 1)
        return owner.strip(), deadline.strip()
    if "/" in value:
        owner, deadline = value.split("/", 1)
        return owner.strip(), deadline.strip()
    return value.strip(), ""


def _simple_topic_number(topic: str) -> str:
    match = re.search(r"Thema\s+0*(\d+)(?:\.(\d+))?", topic or "", re.IGNORECASE)
    if not match:
        return ""
    main = str(int(match.group(1)))
    sub = match.group(2)
    return f"{main}.{sub}" if sub else main


def _simple_verteiler_rows(parsed: ParsedMd) -> list[list[str]]:
    rows = _first_table_from_section(parsed, "Verteiler")
    normalized: list[list[str]] = []
    for row in _clean_rows(rows[1:] if rows else []):
        padded = _pad_row(row, 4)
        if len(row) == 3:
            normalized.append([padded[0], padded[1], "", padded[2]])
        else:
            normalized.append(padded[:4])
    return normalized


def _render_simple_excel_template(parsed: ParsedMd, out_path: Path) -> None:
    if not PROTOKOLL_EINFACH_EXCEL_TEMPLATE.exists():
        raise FileNotFoundError(f"Missing QMG template: {PROTOKOLL_EINFACH_EXCEL_TEMPLATE}")

    wb = load_workbook(str(PROTOKOLL_EINFACH_EXCEL_TEMPLATE))
    header = _simple_header(parsed)

    deck = wb["Deckblatt"]
    _excel_set(deck["A2"], header["title"])
    _excel_set(deck["A3"], header["subtitle"])
    _excel_set(deck["B6"], header["subtitle"])
    _excel_set(deck["B7"], header["project_number"])
    _excel_set(deck["B8"], header["project_name"])
    if header["notice"]:
        _excel_set(deck["C7"], header["notice"])
    _excel_set(deck["A12"], header["ort"])
    _excel_set(deck["B12"], header["gespraechsdatum"])
    _excel_set(deck["C12"], header["erstelldatum"])
    _excel_set(deck["D12"], header["ersteller"])

    teilnehmer = _first_table_from_section(parsed, "Teilnehmer")
    teilnehmer_rows = _clean_rows(teilnehmer[1:] if teilnehmer else [])
    participant_extra = _excel_fill_block(
        deck,
        start_row=16,
        capacity=10,
        rows=[_pad_row(row, 4) for row in teilnehmer_rows],
        cell_indices=[1, 2, 3, 4],
        template_row=16,
        max_col=4,
    )

    verteiler_start = 29 + participant_extra
    verteiler_extra = _excel_fill_block(
        deck,
        start_row=verteiler_start,
        capacity=2,
        rows=_simple_verteiler_rows(parsed),
        cell_indices=[1, 2, 3, 4],
        template_row=verteiler_start,
        max_col=4,
    )
    _set_print_area(deck, "D", 30 + participant_extra + verteiler_extra)

    protocol = wb["Protokoll"]
    inhalt = _first_table_from_section(parsed, "Gesprächsinhalt")
    inhalt_rows = _clean_rows(inhalt[1:] if inhalt else [])
    topic_capacity = 5
    topic_extra = _excel_ensure_capacity(
        protocol,
        start_row=2,
        capacity=topic_capacity,
        needed=len(inhalt_rows),
        template_row=3,
        max_col=5,
    )
    total_topic_rows = max(topic_capacity + topic_extra, len(inhalt_rows))
    for offset in range(total_topic_rows):
        row_idx = 2 + offset
        if offset >= len(inhalt_rows):
            _excel_clear_row(protocol, row_idx, max_col=5)
            continue
        topic, description, owner_deadline = _pad_row(inhalt_rows[offset], 3)
        owner, deadline = _split_owner_deadline(owner_deadline)
        template_row = 2 if re.fullmatch(r"Thema\s+\d+", topic.strip()) else 3
        _excel_put_row(
            protocol,
            row_idx,
            [_simple_topic_number(topic), topic, description, owner, deadline],
            [1, 2, 3, 4, 5],
            template_row=template_row,
            max_col=5,
        )
    _set_print_area(protocol, "E", 6 + topic_extra)

    doku = wb["Doku_Info"]
    termine_section = _section_by_prefix(parsed, "Termine")
    termine: list[list[str]] = []
    if termine_section is not None:
        tables = _tables_from_lines(termine_section.lines)
        termine = _clean_rows(tables[0][1:] if tables else [])
    termine_extra = _excel_fill_block(
        doku,
        start_row=3,
        capacity=3,
        rows=termine,
        cell_indices=[1, 4, 6, 8, 9],
        template_row=3,
        max_col=9,
    )
    doku_shift = termine_extra
    _excel_set(doku.cell(8 + doku_shift, 4), header["ersteller"])
    _excel_set(doku.cell(8 + doku_shift, 8), header["erstelldatum"])
    _excel_set(doku.cell(9 + doku_shift, 4), "")
    _excel_set(doku.cell(9 + doku_shift, 8), "")

    anlagen = _first_table_from_section(parsed, "Anlagen")
    anlagen_rows = _clean_rows(anlagen[1:] if anlagen else [])
    anlagen_start = 17 + doku_shift
    anlagen_extra = _excel_fill_block(
        doku,
        start_row=anlagen_start,
        capacity=3,
        rows=anlagen_rows,
        cell_indices=[1, 5, 6, 8],
        template_row=anlagen_start,
        max_col=9,
    )
    _set_print_area(doku, "I", 32 + doku_shift + anlagen_extra)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(out_path))


def _ensure_protocol_hide_column(ws, last_row: int) -> None:
    _excel_set(ws["H2"], "ausblenden")
    for row_idx in range(3, last_row + 1):
        cell = ws.cell(row_idx, 8)
        _excel_copy_cell(ws["G5"], cell)
        _excel_set_raw(
            cell,
            f'=IF(AND((1+B{row_idx})<Deckblatt!$A$3,G{row_idx}="E"),"x","-")',
        )

    if "Protokoll" in ws.tables:
        table = ws.tables["Protokoll"]
        table.ref = f"A2:H{last_row}"
        existing = list(table.tableColumns)
        if not any(col.name == "ausblenden" for col in existing):
            table.tableColumns.append(TableColumn(id=len(existing) + 1, name="ausblenden"))


def _render_tracking_excel_template(parsed: ParsedMd, out_path: Path) -> None:
    if not TRACKING_EXCEL_TEMPLATE.exists():
        raise FileNotFoundError(f"Missing QMG template: {TRACKING_EXCEL_TEMPLATE}")

    wb = load_workbook(str(TRACKING_EXCEL_TEMPLATE))
    header = _tracking_header(parsed)

    deck = wb["Deckblatt"]
    _excel_set_raw(deck["A3"], _tracking_meeting_number_value(header["meeting_no"]))
    _excel_set(deck["B6"], header["meeting_topic"])
    _excel_set(deck["B7"], header["project_number"])
    _excel_set(deck["B8"], header["project_name"])
    if header["notice"]:
        _excel_set(deck["D7"], header["notice"])
    _excel_set(deck["B10"], header["ort"])
    _excel_set(deck["E10"], header["datum"])
    _excel_set(deck["E11"], header["zeit"])

    teilnehmer = _tracking_data_rows(parsed, "Teilnehmer")
    participant_extra = _excel_fill_block(
        deck,
        start_row=15,
        capacity=10,
        rows=teilnehmer,
        cell_indices=[1, 2, 3, 4, 5, 6, 7],
        template_row=15,
        max_col=7,
    )

    unterlagen = _tracking_data_rows(parsed, "Besprochene Unterlagen")
    underlagen_start = 29 + participant_extra
    underlagen_extra = _excel_fill_block(
        deck,
        start_row=underlagen_start,
        capacity=3,
        rows=[_pad_row(row, 5) for row in unterlagen],
        cell_indices=[1, 4, 5, 6, 7],
        template_row=underlagen_start,
        max_col=7,
    )
    _set_print_area(deck, "G", 32 + participant_extra + underlagen_extra)

    protocol = wb["Protokoll"]
    themen = _tracking_data_rows(parsed, "Besprechungsthemen")
    topic_capacity = 18
    topic_extra = _excel_ensure_capacity(
        protocol,
        start_row=4,
        capacity=topic_capacity,
        needed=len(themen),
        template_row=5,
        max_col=7,
    )
    total_topic_rows = max(topic_capacity + topic_extra, len(themen))
    for offset in range(total_topic_rows):
        row_idx = 4 + offset
        if offset >= len(themen):
            _excel_clear_row(protocol, row_idx, max_col=7)
            continue
        row = themen[offset] + [""] * 7
        is_category = _is_empty_marker(row[1]) and _is_empty_marker(row[2])
        if is_category:
            _excel_put_row(
                protocol,
                row_idx,
                [row[0], row[3]],
                [1, 2],
                template_row=4,
                max_col=7,
            )
        else:
            _excel_put_row(
                protocol,
                row_idx,
                row[:7],
                [1, 2, 3, 4, 5, 6, 7],
                template_row=5,
                max_col=7,
            )
    protocol_last = 22 + topic_extra
    _ensure_protocol_hide_column(protocol, protocol_last)
    _set_print_area(protocol, "G", protocol_last)

    doku = wb["Doku_Info"]
    termine_section = _section_by_prefix(parsed, "Termine")
    termine: list[list[str]] = []
    if termine_section is not None:
        tables = _tables_from_lines(termine_section.lines)
        termine = _clean_rows(tables[0][1:] if tables else [])
    termine_extra = _excel_fill_block(
        doku,
        start_row=3,
        capacity=3,
        rows=termine,
        cell_indices=[1, 4, 6, 8, 9],
        template_row=3,
        max_col=9,
    )
    doku_shift = termine_extra

    aufstell = _tracking_data_rows(parsed, "Aufstellvermerk zum Dokument")
    aufstell_by_role = {row[0].strip().lower(): row for row in aufstell if row}
    for role, row_idx in [("ersteller", 8 + doku_shift), ("geprüft", 9 + doku_shift)]:
        row = aufstell_by_role.get(role, [])
        _excel_set(doku.cell(row_idx, 4), row[1] if len(row) > 1 else "")
        _excel_set(doku.cell(row_idx, 8), row[2] if len(row) > 2 else "")

    anmerkungen = _tracking_data_rows(parsed, "Nachträgliche Anmerkungen zum Dokument")
    anmerkung = anmerkungen[0] if anmerkungen else []
    _excel_set(doku.cell(12 + doku_shift, 3), anmerkung[0] if len(anmerkung) > 0 else "- - -")
    _excel_set(doku.cell(12 + doku_shift, 8), anmerkung[1] if len(anmerkung) > 1 else "- - -")
    _excel_set(doku.cell(13 + doku_shift, 3), anmerkung[2] if len(anmerkung) > 2 else "- - -")

    anlagen = _tracking_data_rows(parsed, "Anlagen")
    anlagen_start = 17 + doku_shift
    anlagen_extra = _excel_fill_block(
        doku,
        start_row=anlagen_start,
        capacity=3,
        rows=anlagen,
        cell_indices=[1, 5, 6, 8],
        template_row=anlagen_start,
        max_col=9,
    )
    doku_shift += anlagen_extra

    kennz = _tracking_data_rows(parsed, "Kennzeichnungen im Dokument")
    kennz_start = 23 + doku_shift
    kennz_extra = _excel_fill_block(
        doku,
        start_row=kennz_start,
        capacity=3,
        rows=kennz,
        cell_indices=[1, 4, 6],
        template_row=kennz_start,
        max_col=9,
    )
    _set_print_area(doku, "I", 32 + doku_shift + kennz_extra)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(out_path))


def render_to_xlsx(parsed: ParsedMd, out_path: Path) -> None:
    if _is_simple_excel_format(parsed):
        _render_simple_excel_template(parsed, out_path)
    elif _is_tracking_excel_format(parsed):
        _render_tracking_excel_template(parsed, out_path)
    else:
        raise ValueError(f"Unsupported XLSX format: {parsed.detected_format}")


# ─── DOCX rendering ────────────────────────────────────────────────────────


def render_to_docx(parsed: ParsedMd, out_path: Path) -> None:
    """Build an EBA-styled DOCX from the parsed MD.

    The supported EBA formats use the official QMG-024-141 Word table shells so
    header/footer/page numbering and EBA-CI are preserved. Unknown Markdown is
    rejected instead of rendered generically."""
    if parsed.detected_format == "gespraechsnotiz":
        _render_gespraechsnotiz_template(parsed, out_path)
        return
    if parsed.detected_format == "protokoll-einfach":
        _render_protokoll_einfach_template(parsed, out_path)
        return
    if parsed.detected_format in {
        "protokoll-tracking",
        "protokoll-lp1-4",
        "protokoll-lp5",
    }:
        _render_tracking_template(parsed, out_path)
        return

    raise ValueError(
        f"Unsupported protocol format '{parsed.detected_format}'. "
        "Refusing to render without a supported QMG template."
    )


# ─── PDF rendering ─────────────────────────────────────────────────────────


def render_to_pdf(docx_path: Path, pdf_path: Path) -> bool:
    """Best-effort DOCX→PDF conversion. Returns True on success.

    Conversion strategies are tried in this order — first one that succeeds
    wins:

      1. **MS Word COM** (Windows, if Word installed) — best fidelity for
         Windows users. Requires `pywin32`.
      2. **LibreOffice headless** (`soffice --headless`) — cross-platform
         (Win/Mac/Linux). The recommended Windows install is the bundled
         LibreOffice from libreoffice.org.
      3. **macOS Pages** via AppleScript — fallback for macOS development.

    All three are graceful: if the converter isn't available, fall through
    to the next. On Windows, pywin32 is installed automatically if it is
    missing and MS Word is available.
    """
    docx_path = docx_path.resolve()
    pdf_path = pdf_path.resolve()
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        pdf_path.unlink()
    except FileNotFoundError:
        pass
    except OSError as exc:
        sys.stderr.write(f"Could not remove stale PDF before export: {exc}\n")
        return False

    # 1. Word COM on Windows
    if sys.platform == "win32":
        pythoncom = None
        word = None
        doc = None
        try:
            try:
                import pythoncom  # type: ignore[import-not-found]
                import win32com.client  # type: ignore[import-not-found]
            except ImportError:
                if not _bootstrap_package("pywin32>=306"):
                    raise
                import pythoncom  # type: ignore[import-not-found]
                import win32com.client  # type: ignore[import-not-found]

            pythoncom.CoInitialize()
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            try:
                doc = word.Documents.Open(
                    FileName=str(docx_path),
                    ReadOnly=True,
                    AddToRecentFiles=False,
                )
                # 17 == wdExportFormatPDF / wdFormatPDF.
                try:
                    doc.ExportAsFixedFormat(str(pdf_path), ExportFormat=17)
                except Exception:
                    doc.SaveAs(str(pdf_path), FileFormat=17)
            finally:
                if doc is not None:
                    try:
                        doc.Close(SaveChanges=False)
                    except Exception:
                        pass
                if word is not None:
                    try:
                        word.Quit()
                    except Exception:
                        pass
                if pythoncom is not None:
                    pythoncom.CoUninitialize()
            if pdf_path.exists():
                return True
        except Exception as exc:
            sys.stderr.write(f"Word COM export skipped/failed: {exc}\n")

    # 2. LibreOffice headless (Windows, Linux, macOS)
    soffice_candidates = ["soffice", "libreoffice"]
    if sys.platform == "win32":
        # Most common Windows install paths
        soffice_candidates = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ] + soffice_candidates
    for cand in soffice_candidates:
        try:
            r = subprocess.run(
                [
                    cand,
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    str(pdf_path.parent),
                    str(docx_path),
                ],
                capture_output=True,
                text=True,
                timeout=180,
            )
            if r.returncode == 0:
                produced = pdf_path.parent / (docx_path.stem + ".pdf")
                if produced != pdf_path and produced.exists():
                    produced.replace(pdf_path)
                if pdf_path.exists():
                    return True
        except FileNotFoundError:
            continue
        except Exception as exc:
            sys.stderr.write(f"LibreOffice export failed: {exc}\n")

    # 3. macOS Pages — dev-environment fallback only
    if sys.platform == "darwin" and Path("/Applications/Pages.app").exists():
        try:
            docx_posix = json.dumps(str(docx_path))
            pdf_posix = json.dumps(str(pdf_path))
            script = (
                'tell application "Pages"\n'
                "  launch\n"
                "  delay 0.5\n"
                f"  set theDoc to open POSIX file {docx_posix}\n"
                "  delay 1\n"
                f"  export theDoc to POSIX file {pdf_posix} as PDF\n"
                "  try\n"
                "    close theDoc saving no\n"
                "  end try\n"
                "end tell\n"
            )
            r = subprocess.run(
                ["osascript", "-e", script],
                capture_output=True,
                text=True,
                timeout=120,
            )
            if pdf_path.exists():
                return True
            sys.stderr.write(f"Pages export failed: {r.stderr.strip()}\n")
        except Exception as exc:
            sys.stderr.write(f"Pages export error: {exc}\n")

    return False


# ─── CLI ───────────────────────────────────────────────────────────────────


def _remove_markdown_intermediate(md_path: Path) -> bool:
    try:
        md_path.unlink()
    except OSError as exc:
        sys.stderr.write(f"Could not remove Markdown intermediate {md_path}: {exc}\n")
        return False
    return True


def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("md_path", type=Path, help="Markdown file to render")
    ap.add_argument("--format", default=None, help="Force format (overrides auto-detect)")
    ap.add_argument("--no-pdf", action="store_true", help="Skip PDF rendering")
    ap.add_argument(
        "--out-dir",
        type=Path,
        default=None,
        help="Output directory (default: alongside MD)",
    )
    args = ap.parse_args(argv)

    md_path: Path = args.md_path.resolve()
    if not md_path.is_file():
        sys.stderr.write(f"Not a file: {md_path}\n")
        return 2

    parsed = parse_protokoll_md(md_path.read_text(encoding="utf-8"))
    if args.format:
        parsed.detected_format = args.format

    out_dir = (args.out_dir or md_path.parent).resolve()
    out_dir.mkdir(parents=True, exist_ok=True)
    docx_path = out_dir / (md_path.stem + ".docx")
    pdf_path = out_dir / (md_path.stem + ".pdf")
    xlsx_path = out_dir / (md_path.stem + ".xlsx")

    if _is_xlsx_only_format(parsed):
        try:
            render_to_xlsx(parsed, xlsx_path)
        except Exception as exc:
            sys.stderr.write(f"XLSX render failed: {exc}\n")
            return 5

        if not _remove_markdown_intermediate(md_path):
            return 6

        print(f"XLSX: {xlsx_path}")
        print(f"Format: {parsed.detected_format}")
        return 0

    try:
        render_to_docx(parsed, docx_path)
    except Exception as exc:
        sys.stderr.write(f"DOCX render failed: {exc}\n")
        return 3

    pdf_ok = False
    if not args.no_pdf:
        pdf_ok = render_to_pdf(docx_path, pdf_path)

    pdf_required = not args.no_pdf
    if pdf_required and not pdf_ok:
        sys.stderr.write(
            "Required PDF render failed. On Windows, the renderer already attempted "
            "the user-local pywin32 bootstrap and MS Word COM export; ensure Word "
            "can open normally. On macOS/Linux, install or repair one of the "
            "document converters listed in ausgabe-konvention.md, then rerun.\n"
        )
        return 4

    if not _remove_markdown_intermediate(md_path):
        return 6

    print(f"DOCX: {docx_path}")
    if pdf_ok:
        print(f"PDF:  {pdf_path}")
    print(f"Format: {parsed.detected_format}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
