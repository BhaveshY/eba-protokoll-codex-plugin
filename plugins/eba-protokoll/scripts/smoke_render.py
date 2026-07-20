#!/usr/bin/env python3
"""Smoke-test the DOCX renderer against repository examples.

Developer test dependencies:
    python3 -m pip install -r scripts/requirements.txt

The core matrix uses --no-pdf so it remains independent of Word, LibreOffice,
or Pages. When LibreOffice, pdfinfo, and pdftotext are available, an additional
multi-page regression verifies header references and page totals in real PDFs.
End users do not run this; render_protokoll.py bootstraps its own dependencies.
"""

from __future__ import annotations

from collections import Counter
import re
import shutil
import subprocess
import sys
import tempfile
from zipfile import ZipFile
from pathlib import Path
from xml.etree import ElementTree as ET

try:
    from docx import Document
    from openpyxl import load_workbook
except ImportError:
    sys.stderr.write(
        "smoke_render.py: missing dependency 'python-docx' or 'openpyxl'.\n"
        "Install with: python3 -m pip install -r scripts/requirements.txt\n"
    )
    sys.exit(3)


REPO_ROOT = Path(__file__).resolve().parents[1]
RENDERER = REPO_ROOT / "scripts" / "render_protokoll.py"
W_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
R_ID = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"
W_TYPE = f"{{{W_NS['w']}}}type"
DOCX_TEMPLATE_BY_EXAMPLE = {
    "references/examples/beispiel-ausgabe-gespraechsnotiz.md":
        "references/templates/qmg/QMG-024-141_ORG-GESPRAECHSNOTIZ_230202-D.docx",
    "references/examples/beispiel-ausgabe-eba-interview.md":
        "references/templates/qmg/QMG-024-141_ORG-GESPRAECHSNOTIZ_230202-D.docx",
    "references/examples/beispiel-ausgabe-einfach.md":
        "references/templates/qmg/QMG-024-141_ORG-PK-LP1-4-MA_230227-A.docx",
    "references/examples/beispiel-ausgabe-lp1-4.md":
        "references/templates/qmg/QMG-024-141_ORG-PK-LP5-MA_230202-B.docx",
    "references/examples/beispiel-ausgabe-lp5.md":
        "references/templates/qmg/QMG-024-141_ORG-PK-LP5-MA_230202-B.docx",
}
QMG_TEMPLATE_EXAMPLES = set(DOCX_TEMPLATE_BY_EXAMPLE)
HEADER_REF_TARGETS_BY_EXAMPLE = {
    "references/examples/beispiel-ausgabe-gespraechsnotiz.md":
        {"PrjNr", "PrjName", "Name", "DokBeschr"},
    "references/examples/beispiel-ausgabe-eba-interview.md":
        {"PrjNr", "PrjName", "Name", "DokBeschr"},
    "references/examples/beispiel-ausgabe-einfach.md":
        {"PrjNr", "PrjName", "Name", "DokBeschr"},
    "references/examples/beispiel-ausgabe-lp1-4.md":
        {"PrjNr", "PrjName", "Besprechung"},
    "references/examples/beispiel-ausgabe-lp5.md":
        {"PrjNr", "PrjName", "Besprechung"},
}
SIMPLE_XLSX_TEMPLATE = REPO_ROOT / "references/templates/qmg/QMG-024-141_ORG-PK-LP1-4-EXCEL-MA_240920-A.xlsx"
TRACKING_XLSX_TEMPLATE = REPO_ROOT / "references/templates/qmg/QMG-024-141_ORG-PK-EXCEL-MA_240926-C.xlsx"
TRACKING_XLSX_EXAMPLES = {
    "references/examples/beispiel-ausgabe-bim.md",
}
XLSX_ONLY_EXAMPLES = TRACKING_XLSX_EXAMPLES


def output_contract_checks(
    out_dir: Path,
    stdout: str,
    expected_suffixes: set[str],
    label: str,
) -> list[str]:
    failures: list[str] = []
    if not out_dir.is_dir():
        return [f"{label}: output directory was not created"]
    actual_suffixes = {path.suffix for path in out_dir.iterdir() if path.is_file()}
    if actual_suffixes != expected_suffixes:
        failures.append(
            f"{label}: expected only {sorted(expected_suffixes)}, got {sorted(actual_suffixes)}"
        )
    if ".md" in actual_suffixes or "MD:" in stdout or "Markdown:" in stdout:
        failures.append(f"{label}: Markdown leaked into the final output contract")
    expected_labels = {".docx": "DOCX:", ".pdf": "PDF:", ".xlsx": "XLSX:"}
    for suffix, output_label in expected_labels.items():
        if (suffix in expected_suffixes) != (output_label in stdout):
            failures.append(f"{label}: stdout label {output_label} does not match artifacts")
    return failures


def read_docx_text(path: Path) -> str:
    doc = Document(str(path))
    parts: list[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def read_xlsx_text(path: Path) -> str:
    wb = load_workbook(path, data_only=False)
    parts: list[str] = []
    for sheet_name in ["Deckblatt", "Protokoll", "Doku_Info"]:
        ws = wb[sheet_name]
        for row in ws.iter_rows():
            values = [str(cell.value) for cell in row if cell.value is not None]
            if values:
                parts.append(" | ".join(values))
    return "\n".join(parts)


def docx_relationship_targets(z: ZipFile) -> list[tuple[str, str | None]]:
    rels = ET.fromstring(z.read("word/_rels/document.xml.rels"))
    targets: list[tuple[str, str | None]] = []
    for rel in rels:
        rel_type = rel.attrib.get("Type", "")
        if rel_type.endswith(("header", "footer", "styles", "settings", "theme")):
            targets.append((rel_type.rsplit("/", 1)[-1], rel.attrib.get("Target")))
    return sorted(targets)


def docx_style_ids(z: ZipFile) -> set[str]:
    styles = ET.fromstring(z.read("word/styles.xml"))
    return {
        style.attrib.get(f"{{{W_NS['w']}}}styleId", "")
        for style in styles.findall("w:style", W_NS)
    }


def docx_first_section_refs(z: ZipFile) -> list[tuple[str, str, str | None]]:
    document = ET.fromstring(z.read("word/document.xml"))
    sections = document.findall(".//w:sectPr", W_NS)
    if not sections:
        return []
    refs: list[tuple[str, str, str | None]] = []
    for tag, label in [("w:headerReference", "header"), ("w:footerReference", "footer")]:
        for ref in sections[0].findall(tag, W_NS):
            refs.append((label, ref.attrib.get(W_TYPE, "default"), ref.attrib.get(R_ID)))
    return sorted(refs)


def xml_text(z: ZipFile, part: str, namespace: str, tag: str) -> str:
    root = ET.fromstring(z.read(part))
    node = root.find(f"{{{namespace}}}{tag}")
    return node.text if node is not None and node.text is not None else ""


def docx_core_metadata(z: ZipFile) -> dict[str, str]:
    return {
        "creator": xml_text(z, "docProps/core.xml", "http://purl.org/dc/elements/1.1/", "creator"),
        "lastModifiedBy": xml_text(
            z,
            "docProps/core.xml",
            "http://schemas.openxmlformats.org/package/2006/metadata/core-properties",
            "lastModifiedBy",
        ),
        "category": xml_text(
            z,
            "docProps/core.xml",
            "http://schemas.openxmlformats.org/package/2006/metadata/core-properties",
            "category",
        ),
        "company": xml_text(
            z,
            "docProps/app.xml",
            "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties",
            "Company",
        ),
        "application": xml_text(
            z,
            "docProps/app.xml",
            "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties",
            "Application",
        ),
    }


def docx_bookmark_boundaries(z: ZipFile) -> tuple[list[tuple[str, str]], list[str]]:
    document = ET.fromstring(z.read("word/document.xml"))
    starts = [
        (
            node.attrib.get(f"{{{W_NS['w']}}}name", ""),
            node.attrib.get(f"{{{W_NS['w']}}}id", ""),
        )
        for node in document.findall(".//w:bookmarkStart", W_NS)
    ]
    ends = [
        node.attrib.get(f"{{{W_NS['w']}}}id", "")
        for node in document.findall(".//w:bookmarkEnd", W_NS)
    ]
    return starts, ends


def docx_header_ref_targets(z: ZipFile) -> set[str]:
    targets: set[str] = set()
    for name in z.namelist():
        if not name.startswith("word/header") or not name.endswith(".xml"):
            continue
        root = ET.fromstring(z.read(name))
        for instruction in root.findall(".//w:instrText", W_NS):
            match = re.search(r"\bREF\s+([A-Za-z0-9_]+)", instruction.text or "")
            if match:
                targets.add(match.group(1))
    return targets


def _column_widths(ws) -> dict[str, float | None]:
    return {key: dim.width for key, dim in ws.column_dimensions.items()}


def _row_heights(ws) -> dict[int, float | None]:
    return {key: dim.height for key, dim in ws.row_dimensions.items()}


def _merged_ranges(ws) -> list[str]:
    return sorted(str(rng) for rng in ws.merged_cells.ranges)


def _sheet_values(ws) -> list[tuple[int, int, object]]:
    values: list[tuple[int, int, object]] = []
    for row in ws.iter_rows():
        for cell in row:
            if cell.value is not None:
                values.append((cell.row, cell.column, cell.value))
    return values


def xlsx_official_template_checks(path: Path, template_path: Path, example: str) -> list[str]:
    failures: list[str] = []
    wb = load_workbook(path, data_only=False)
    template = load_workbook(template_path, data_only=False)
    if wb.sheetnames != template.sheetnames:
        failures.append(
            f"{example}: workbook sheets no longer match official template: {wb.sheetnames}"
        )
    if wb.properties.creator != template.properties.creator:
        failures.append(f"{example}: workbook creator metadata changed")
    if wb.properties.lastModifiedBy != template.properties.lastModifiedBy:
        failures.append(f"{example}: workbook lastModifiedBy metadata changed")
    if [style.name for style in wb._named_styles] != [style.name for style in template._named_styles]:
        failures.append(f"{example}: workbook named styles changed")

    for sheet_name in template.sheetnames:
        if sheet_name not in wb.sheetnames:
            continue
        ws = wb[sheet_name]
        template_ws = template[sheet_name]
        if ws.sheet_state != template_ws.sheet_state:
            failures.append(f"{example}: {sheet_name} visibility changed")
        if ws.freeze_panes != template_ws.freeze_panes:
            failures.append(f"{example}: {sheet_name} freeze panes changed")
        if _column_widths(ws) != _column_widths(template_ws):
            failures.append(f"{example}: {sheet_name} column widths changed")

    for sheet_name in ["Hilfe und Tipps", "intern"]:
        ws = wb[sheet_name]
        template_ws = template[sheet_name]
        if (ws.max_row, ws.max_column) != (template_ws.max_row, template_ws.max_column):
            failures.append(f"{example}: {sheet_name} dimensions changed")
        if ws.print_area != template_ws.print_area:
            failures.append(f"{example}: {sheet_name} print area changed")
        if _row_heights(ws) != _row_heights(template_ws):
            failures.append(f"{example}: {sheet_name} row heights changed")
        if _merged_ranges(ws) != _merged_ranges(template_ws):
            failures.append(f"{example}: {sheet_name} merged cells changed")
        if _sheet_values(ws) != _sheet_values(template_ws):
            failures.append(f"{example}: {sheet_name} content changed")
    return failures


def xlsx_template_checks(path: Path, example: str) -> list[str]:
    failures: list[str] = []
    wb = load_workbook(path, data_only=False)
    failures.extend(xlsx_official_template_checks(path, TRACKING_XLSX_TEMPLATE, example))
    expected_sheets = ["Deckblatt", "Protokoll", "Doku_Info", "Hilfe und Tipps", "intern"]
    if wb.sheetnames != expected_sheets:
        failures.append(f"{example}: XLSX sheet structure changed: {wb.sheetnames}")
    if "Protokoll" not in wb["Protokoll"].tables:
        failures.append(f"{example}: XLSX Protokoll table is missing")
    else:
        ref = wb["Protokoll"].tables["Protokoll"].ref
        if not ref.startswith("A2:H"):
            failures.append(f"{example}: XLSX Protokoll table should include ausblenden column H (got {ref})")
    if wb["Protokoll"]["H2"].value != "ausblenden":
        failures.append(f"{example}: XLSX ausblenden helper header missing")
    if not str(wb["Protokoll"]["H5"].value or "").startswith("=IF(AND((1+B5)<Deckblatt!$A$3"):
        failures.append(f"{example}: XLSX ausblenden helper formula missing")
    if not isinstance(wb["Deckblatt"]["A3"].value, int):
        failures.append(f"{example}: XLSX meeting number should be numeric in Deckblatt!A3")
    for sheet_name in ["Deckblatt", "Protokoll", "Doku_Info"]:
        sheet_text = "\n".join(
            str(cell.value)
            for row in wb[sheet_name].iter_rows()
            for cell in row
            if cell.value is not None
        )
        for placeholder in [
            "_Vorname_",
            "_Name_",
            "_Firma_",
            "_ Dokument/e, Plan/Pläne _",
            "_Thema 01_",
            "_Ersteller eintragen_",
            "Besprechnungsthema A",
        ]:
            if placeholder in sheet_text:
                failures.append(f"{example}: XLSX placeholder leaked into {sheet_name}: {placeholder}")
    return failures


def simple_xlsx_template_checks(path: Path, example: str) -> list[str]:
    failures: list[str] = []
    wb = load_workbook(path, data_only=False)
    failures.extend(xlsx_official_template_checks(path, SIMPLE_XLSX_TEMPLATE, example))
    expected_sheets = ["Deckblatt", "Protokoll", "Doku_Info", "Hilfe und Tipps", "intern"]
    if wb.sheetnames != expected_sheets:
        failures.append(f"{example}: simple XLSX sheet structure changed: {wb.sheetnames}")
    if wb["Protokoll"]["A1"].value != "Gesprächsinhalt":
        failures.append(f"{example}: simple XLSX Protokoll sheet is not the official simple template")
    if wb["Protokoll"]["D1"].value != "zuständig" or wb["Protokoll"]["E1"].value != "Frist":
        failures.append(f"{example}: simple XLSX zuständig/Frist headers missing")
    if wb["Protokoll"]["A2"].value != "1" or wb["Protokoll"]["B3"].value != "Thema 01.1":
        failures.append(f"{example}: simple XLSX topic numbering not populated")
    if "Protokoll" in wb["Protokoll"].tables:
        failures.append(f"{example}: simple XLSX should not use the D/K tracking table")
    if wb["Protokoll"]["H2"].value == "ausblenden":
        failures.append(f"{example}: simple XLSX should not include tracking helper column H")
    for sheet_name in ["Deckblatt", "Protokoll", "Doku_Info"]:
        sheet_text = "\n".join(
            str(cell.value)
            for row in wb[sheet_name].iter_rows()
            for cell in row
            if cell.value is not None
        )
        for placeholder in [
            "_Vorname_",
            "_Name_",
            "_Firma_",
            "_Kürzel_",
            "_Prj.-Nr._",
            "_Prj.-Name_",
            "_Besprechungsthema_",
            "_Ersteller_",
            "_Ersteller eintragen_",
            "_ Dokument/e, Plan/Pläne _",
            "_Thema 01_",
            "_Dokument/e, Plan/Pläne_",
            "_Format_",
            "Besprechnungsthema A",
        ]:
            if placeholder in sheet_text:
                failures.append(f"{example}: simple XLSX placeholder leaked into {sheet_name}: {placeholder}")
    return failures


def qmg_template_checks(path: Path, example: str) -> list[str]:
    if example not in QMG_TEMPLATE_EXAMPLES:
        return []
    failures: list[str] = []
    template_path = REPO_ROOT / DOCX_TEMPLATE_BY_EXAMPLE[example]
    with ZipFile(path) as z, ZipFile(template_path) as template_z:
        names = z.namelist()
        body = ET.fromstring(z.read("word/document.xml")).find("w:body", W_NS)
        if body is None:
            return [f"{example}: DOCX has no body"]
        body_tables = body.findall("w:tbl", W_NS)
        if len(body_tables) != 4:
            failures.append(f"{example}: rendered output should keep exactly 4 QMG body tables")
        xml_blob = "\n".join(
            z.read(name).decode("utf-8", "ignore")
            for name in names
            if name.startswith("word/") and name.endswith(".xml")
        )
        for internal_marker in ["Hilfe und Tipps", "Dokument-Raster", "Diese Zeile bitte nicht löschen"]:
            if internal_marker in xml_blob:
                failures.append(f"{example}: internal QMG helper page leaked into output")
        for placeholder in [
            "_Vorname_",
            "_Name_",
            "_Firma_",
            "_ Dokument/e, Plan/Pläne _",
            "_Thema 01_",
            "_Beschreibung einfügen_",
            "Besprechnungsthema A",
        ]:
            if placeholder in xml_blob:
                failures.append(f"{example}: QMG placeholder leaked into output: {placeholder}")
        headers = [name for name in names if name.startswith("word/header")]
        footers = [name for name in names if name.startswith("word/footer")]
        if not headers:
            failures.append(f"{example}: official header parts are missing")
        if not footers:
            failures.append(f"{example}: official footer parts are missing")
        header_text = "\n".join(z.read(name).decode("utf-8", "ignore") for name in headers)
        footer_text = "\n".join(z.read(name).decode("utf-8", "ignore") for name in footers)
        if "Eike Becker_Architekten" not in header_text and "Eike Becker_Architekten" not in footer_text:
            failures.append(f"{example}: EBA header/footer branding missing")
        if "PAGE" not in footer_text or "NUMPAGES" not in footer_text:
            failures.append(f"{example}: page number fields missing from footer")
        if "SECTIONPAGES" in footer_text:
            failures.append(f"{example}: stale source-template SECTIONPAGES field remains")
        if re.search(r"REF\s+\\\*\s+CHARFORMAT\s+PrjNr", header_text):
            failures.append(f"{example}: malformed source-template PrjNr REF field remains")

        starts, ends = docx_bookmark_boundaries(z)
        bookmark_names = [name for name, _ in starts]
        start_ids = [bookmark_id for _, bookmark_id in starts]
        expected_targets = HEADER_REF_TARGETS_BY_EXAMPLE[example]
        actual_targets = docx_header_ref_targets(z)
        if actual_targets != expected_targets:
            failures.append(
                f"{example}: expected header REF targets {sorted(expected_targets)}, "
                f"got {sorted(actual_targets)}"
            )
        missing_bookmarks = expected_targets - set(bookmark_names)
        if missing_bookmarks:
            failures.append(
                f"{example}: header REF bookmark targets missing: {sorted(missing_bookmarks)}"
            )
        if any(count != 1 for count in Counter(bookmark_names).values()):
            failures.append(f"{example}: duplicate bookmark names in rendered DOCX")
        if any(count != 1 for count in Counter(start_ids).values()):
            failures.append(f"{example}: duplicate bookmark start IDs in rendered DOCX")
        if Counter(start_ids) != Counter(ends):
            failures.append(f"{example}: bookmark start/end boundaries do not match")
        settings_text = z.read("word/settings.xml").decode("utf-8", "ignore")
        if "updateFields" not in settings_text:
            failures.append(f"{example}: rendered DOCX does not request field refresh")

        if docx_relationship_targets(z) != docx_relationship_targets(template_z):
            failures.append(f"{example}: DOCX template header/footer/style relationships changed")
        if z.read("word/theme/theme1.xml") != template_z.read("word/theme/theme1.xml"):
            failures.append(f"{example}: DOCX theme part no longer matches the official template")
        if docx_style_ids(z) != docx_style_ids(template_z):
            failures.append(f"{example}: DOCX styles no longer match the official template")
        if docx_first_section_refs(z) != docx_first_section_refs(template_z):
            failures.append(f"{example}: DOCX first-section header/footer references changed")
        if docx_core_metadata(z) != docx_core_metadata(template_z):
            failures.append(f"{example}: DOCX core/app metadata changed")
        rendered_custom_xml = sorted(name for name in names if name.startswith("customXml/"))
        template_custom_xml = sorted(name for name in template_z.namelist() if name.startswith("customXml/"))
        if rendered_custom_xml != template_custom_xml:
            failures.append(f"{example}: DOCX custom XML metadata parts changed")
    return failures


def render_example(example: str, required_text: list[str]) -> list[str]:
    failures: list[str] = []
    src = REPO_ROOT / example
    with tempfile.TemporaryDirectory(prefix="eba-render-smoke-") as tmp:
        tmp_path = Path(tmp)
        md_path = tmp_path / src.name
        out_dir = tmp_path / "out"
        shutil.copy2(src, md_path)
        result = subprocess.run(
            [
                sys.executable,
                str(RENDERER),
                str(md_path),
                "--out-dir",
                str(out_dir),
                "--no-pdf",
            ],
            cwd=str(REPO_ROOT),
            text=True,
            capture_output=True,
            timeout=120,
        )
        if result.returncode != 0:
            failures.append(
                f"{example}: renderer exited {result.returncode}: {result.stderr.strip()}"
            )
            return failures
        if md_path.exists():
            failures.append(f"{example}: markdown intermediate was not deleted")

        docx_path = out_dir / f"{md_path.stem}.docx"
        xlsx_path = out_dir / f"{md_path.stem}.xlsx"
        if example in XLSX_ONLY_EXAMPLES:
            failures.extend(output_contract_checks(out_dir, result.stdout, {".xlsx"}, example))
            if docx_path.exists():
                failures.append(f"{example}: Excel-origin format unexpectedly wrote DOCX")
            if not xlsx_path.exists():
                failures.append(f"{example}: XLSX was not written")
                return failures
            rendered_xlsx = read_xlsx_text(xlsx_path)
            for needle in [
                "BIM-Koordination JF-07",
                "BIMcollab-Issue-Liste",
                "FusionLive bleibt die verbindliche CDE",
                "erledigt",
            ]:
                if needle not in rendered_xlsx:
                    failures.append(f"{example}: rendered XLSX missing {needle!r}")
            failures.extend(xlsx_template_checks(xlsx_path, example))
            return failures

        failures.extend(output_contract_checks(out_dir, result.stdout, {".docx"}, example))
        if not docx_path.exists():
            failures.append(f"{example}: DOCX was not written")
            return failures
        rendered = read_docx_text(docx_path)
        for needle in required_text:
            if needle not in rendered:
                failures.append(f"{example}: rendered DOCX missing {needle!r}")
        failures.extend(qmg_template_checks(docx_path, example))

        if xlsx_path.exists():
            failures.append(f"{example}: unexpected XLSX was written for non-Excel format")
    return failures


def forced_excel_formats() -> list[str]:
    failures: list[str] = []
    cases = [
        (
            "references/examples/beispiel-ausgabe-einfach.md",
            "protokoll-einfach-excel",
            simple_xlsx_template_checks,
            [
                "Kick-Off Meeting Projekt VTS-549",
                "Projektorganisation: EBA übernimmt die Gesamtkoordination.",
                "22.04.26",
            ],
        ),
        (
            "references/examples/beispiel-ausgabe-lp1-4.md",
            "protokoll-lp1-4-excel",
            xlsx_template_checks,
            [
                "Planungsbesprechung — BIM, Bauantrag, Wohnfassade",
                "LP3-Modell (FusionLive-Upload)",
                "DGNB-Workshop am 14.04.26",
            ],
        ),
    ]
    for example, forced_format, check_fn, required_text in cases:
        src = REPO_ROOT / example
        with tempfile.TemporaryDirectory(prefix="eba-render-forced-xlsx-") as tmp:
            tmp_path = Path(tmp)
            md_path = tmp_path / src.name
            out_dir = tmp_path / "out"
            shutil.copy2(src, md_path)
            result = subprocess.run(
                [
                    sys.executable,
                    str(RENDERER),
                    str(md_path),
                    "--format",
                    forced_format,
                    "--out-dir",
                    str(out_dir),
                    "--no-pdf",
                ],
                cwd=str(REPO_ROOT),
                text=True,
                capture_output=True,
                timeout=120,
            )
            if result.returncode != 0:
                failures.append(
                    f"{example} forced {forced_format}: renderer exited "
                    f"{result.returncode}: {result.stderr.strip()}"
                )
                continue
            if md_path.exists():
                failures.append(
                    f"{example} forced {forced_format}: markdown intermediate was not deleted"
                )
            failures.extend(
                output_contract_checks(
                    out_dir,
                    result.stdout,
                    {".xlsx"},
                    f"{example} forced {forced_format}",
                )
            )
            docx_path = out_dir / f"{md_path.stem}.docx"
            xlsx_path = out_dir / f"{md_path.stem}.xlsx"
            if docx_path.exists():
                failures.append(f"{example} forced {forced_format}: unexpected DOCX was written")
            if not xlsx_path.exists():
                failures.append(f"{example} forced {forced_format}: XLSX was not written")
                continue
            rendered_xlsx = read_xlsx_text(xlsx_path)
            for needle in required_text:
                if needle not in rendered_xlsx:
                    failures.append(
                        f"{example} forced {forced_format}: rendered XLSX missing {needle!r}"
                    )
            failures.extend(check_fn(xlsx_path, f"{example} forced {forced_format}"))
    return failures


def unknown_format_rejected() -> list[str]:
    failures: list[str] = []
    with tempfile.TemporaryDirectory(prefix="eba-render-unknown-") as tmp:
        tmp_path = Path(tmp)
        md_path = tmp_path / "unknown.md"
        out_dir = tmp_path / "out"
        md_path.write_text("# Freitext\n\nDies ist kein EBA-Protokoll.\n", encoding="utf-8")
        result = subprocess.run(
            [
                sys.executable,
                str(RENDERER),
                str(md_path),
                "--out-dir",
                str(out_dir),
                "--no-pdf",
            ],
            cwd=str(REPO_ROOT),
            text=True,
            capture_output=True,
            timeout=120,
        )
        if result.returncode != 3:
            failures.append(f"unknown format: expected renderer exit 3, got {result.returncode}")
        if "Refusing to render without a supported QMG template" not in result.stderr:
            failures.append("unknown format: renderer did not explain QMG-template refusal")
        if (out_dir / "unknown.docx").exists():
            failures.append("unknown format: generic DOCX was written")
    return failures


def paths_with_spaces_supported() -> list[str]:
    failures: list[str] = []
    src = REPO_ROOT / "references/examples/beispiel-ausgabe-einfach.md"
    with tempfile.TemporaryDirectory(prefix="eba render smoke spaces ") as tmp:
        tmp_path = Path(tmp)
        md_path = tmp_path / "source with spaces.md"
        out_dir = tmp_path / "out with spaces"
        shutil.copy2(src, md_path)
        result = subprocess.run(
            [
                sys.executable,
                str(RENDERER),
                str(md_path),
                "--format",
                "protokoll-einfach",
                "--out-dir",
                str(out_dir),
                "--no-pdf",
            ],
            cwd=str(REPO_ROOT),
            text=True,
            capture_output=True,
            timeout=120,
        )
        if result.returncode != 0:
            failures.append(
                f"paths with spaces: renderer exited {result.returncode}: {result.stderr.strip()}"
            )
        if not (out_dir / "source with spaces.docx").exists():
            failures.append("paths with spaces: DOCX was not written")
        if md_path.exists():
            failures.append("paths with spaces: markdown intermediate was not deleted")
        failures.extend(
            output_contract_checks(out_dir, result.stdout, {".docx"}, "paths with spaces")
        )
    return failures


def removed_debug_flags_rejected() -> list[str]:
    failures: list[str] = []
    for flag in ["--keep-md", "--no-xlsx", "--pdf-optional"]:
        result = subprocess.run(
            [sys.executable, str(RENDERER), "placeholder.md", flag],
            cwd=str(REPO_ROOT),
            text=True,
            capture_output=True,
            timeout=30,
        )
        if result.returncode != 2 or "unrecognized arguments" not in result.stderr:
            failures.append(f"removed flag {flag}: renderer still accepts the dead debug path")
    return failures


def libreoffice_pdf_regression_checks() -> list[str]:
    """Exercise the portable PDF path when its local toolchain is available."""
    converter = shutil.which("soffice") or shutil.which("libreoffice")
    pdftotext = shutil.which("pdftotext")
    pdfinfo = shutil.which("pdfinfo")
    if not converter or not pdftotext or not pdfinfo:
        return []

    failures: list[str] = []
    cases = [
        "references/examples/beispiel-ausgabe-einfach.md",
        "references/examples/beispiel-ausgabe-lp1-4.md",
    ]
    for example in cases:
        src = REPO_ROOT / example
        with tempfile.TemporaryDirectory(prefix="eba-render-pdf-regression-") as tmp:
            tmp_path = Path(tmp)
            md_path = tmp_path / src.name
            out_dir = tmp_path / "out"
            shutil.copy2(src, md_path)
            result = subprocess.run(
                [
                    sys.executable,
                    str(RENDERER),
                    str(md_path),
                    "--out-dir",
                    str(out_dir),
                ],
                cwd=str(REPO_ROOT),
                text=True,
                capture_output=True,
                timeout=240,
            )
            label = f"{example} LibreOffice PDF"
            if result.returncode != 0:
                failures.append(
                    f"{label}: renderer exited {result.returncode}: {result.stderr.strip()}"
                )
                continue
            failures.extend(
                output_contract_checks(out_dir, result.stdout, {".docx", ".pdf"}, label)
            )
            pdf_path = out_dir / f"{md_path.stem}.pdf"
            if not pdf_path.exists():
                failures.append(f"{label}: PDF was not written")
                continue
            info = subprocess.run(
                [pdfinfo, str(pdf_path)],
                text=True,
                capture_output=True,
                timeout=30,
                check=False,
            )
            page_match = re.search(r"^Pages:\s+(\d+)$", info.stdout, re.MULTILINE)
            if info.returncode != 0 or not page_match:
                failures.append(f"{label}: could not determine PDF page count")
                continue
            page_count = int(page_match.group(1))
            if page_count < 2:
                failures.append(f"{label}: regression fixture did not produce multiple pages")

            extracted = subprocess.run(
                [pdftotext, "-layout", str(pdf_path), "-"],
                text=True,
                capture_output=True,
                timeout=30,
                check=False,
            )
            if extracted.returncode != 0:
                failures.append(f"{label}: pdftotext failed")
                continue
            pdf_text = extracted.stdout
            if "Reference source not found" in pdf_text or "Error:" in pdf_text:
                failures.append(f"{label}: broken template field is visible in PDF")
            for page_number in range(1, page_count + 1):
                marker = f"{page_number}_{page_count}"
                if marker not in pdf_text:
                    failures.append(f"{label}: expected page marker {marker!r} is missing")
    return failures


def main() -> int:
    checks = {
        "references/examples/beispiel-ausgabe-gespraechsnotiz.md": [
            "Gesprächsnotiz",
            "Bauantragsstand und Rückmeldung der Bauaufsicht",
            "Werden innerhalb von 3 Kalendertagen",
        ],
        "references/examples/beispiel-ausgabe-einfach.md": [
            "Protokoll",
            "Kick-Off Meeting Projekt VTS-549",
            "Zuständig/Frist",
        ],
        "references/examples/beispiel-ausgabe-lp1-4.md": [
            "zur Besprechung Nr. 12",
            "Planungsbesprechung — BIM, Bauantrag, Wohnfassade",
            "LP3-Modell (FusionLive-Upload)",
            "Ort | Online",
            "Datum | 24.03.26",
            "Zeit | 09:00 – 09:07",
            "Werden innerhalb von 5 Kalendertagen",
            "Besprechungsthemen",
        ],
        "references/examples/beispiel-ausgabe-lp5.md": [
            "zur Besprechung Nr. 8",
            "Baubesprechung — Rohbau, Mängel, Brandschutz",
            "Witterung",
            "Schalungsplan UG1, UG2",
            "M-048",
        ],
        "references/examples/beispiel-ausgabe-bim.md": [
            "zur Besprechung Nr. 07",
            "BIM-Koordination JF-07",
            "BIMcollab-Issue-Liste",
            "Ort | Online",
            "D/K | B | LN",
        ],
    }

    failures: list[str] = []
    for example, required_text in checks.items():
        failures.extend(render_example(example, required_text))
    failures.extend(forced_excel_formats())
    failures.extend(unknown_format_rejected())
    failures.extend(paths_with_spaces_supported())
    failures.extend(removed_debug_flags_rejected())
    failures.extend(libreoffice_pdf_regression_checks())

    if failures:
        print(f"Render smoke test failed with {len(failures)} issue(s):")
        for failure in failures:
            print(f"- {failure}")
        return 1
    print("Render smoke test passed.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
